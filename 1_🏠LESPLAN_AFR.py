import streamlit as st
import os
import io
import docx
from datetime import date
from docx.shared import Pt, RGBColor
from pathlib import Path
from PIL import Image
from streamlit_extras.colored_header import colored_header
from streamlit_extras.app_logo import add_logo

# Define paths
THIS_DIR = Path(__file__).parent if "__file__" in locals() else Path.cwd()
STYLES_DIR = THIS_DIR / "styles"
CSS_FILE = STYLES_DIR / "main.css"
IMAGES_DIR = THIS_DIR / "IMAGES"
HEADER_IMAGE = IMAGES_DIR / "header.png"
LOGO_IMAGE = IMAGES_DIR / "wapen.png"

# Load CSS file
def load_css_file(css_file_path):
    try:
        with open(css_file_path) as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    except FileNotFoundError:
        st.error(f"CSS file not found at {css_file_path}. Please ensure the file exists.")

# Set page title and icon
st.set_page_config(page_title="Lesson Plan Creator", page_icon=":books:", layout="wide")

# Load CSS
load_css_file(CSS_FILE)

# Display header image
try:
    st.image(str(HEADER_IMAGE))
except FileNotFoundError:
    st.error(f"Header image not found at {HEADER_IMAGE}. Please ensure the file exists.")

# Website redirect notice
st.header("This website is moving to a new domain soon and will be deleted:")
KLIK_HIER = "https://resourceshssd.streamlit.app"  # Define the URL
st.markdown(
    f'<a href="{KLIK_HIER}" class="button">👉 KLIK HIER</a>',
    unsafe_allow_html=True,
)
st.subheader("Die nuwe site kan gevind word by resourceshssd.streamlit.app")
st.write("Groete")
st.write("Mr. Visagie @ Saul Damon High School")

# Add app logo
try:
    add_logo(str(LOGO_IMAGE), height=150)
except FileNotFoundError:
    st.error(f"Logo image not found at {LOGO_IMAGE}. Please ensure the file exists.")

# Colored header
colored_header(
    label="LESBEPLANNER",
    description="Hierdie interaktiewe webblad help u om maklik en vinnig lesplanne te skep.",
    color_name="red-70",
)

# Create input fields
st.write("Vul asseblief die volgende velde in om 'n nuwe lesplan te skep:")
st.write("")
subject = st.text_input("VAK")
lesson_title = st.text_input("LES TITEL")
grade = st.selectbox("GRAAD", ["9", "10", "11", "12"])
st.subheader("KIES JOU BEGIN- EN EINDDATUM")
start_date = st.date_input("VANAF", value=date.today())
end_date = st.date_input("TOT", value=date.today())
lesson_objective = st.text_area("LES DOELWITTE")
lesson_introduction = st.text_area("INLEIDING")
lesson_activities = st.text_area("LES AKTIWITEITE")
materials_needed = st.text_area("MATERIAAL BENODIG")
homework = st.text_area("HUISWERK")
notes = st.text_area("NOTAS")
st.subheader("OPVOEDER BESONDERHEDE")
teacher_name = st.text_input("VOORLETTERS")
teacher_surname = st.text_input("VAN")
st.write("")
st.write("Created by Mr. A.R Visagie @ Saul Damon High School")

# Create save button
if st.button("Create Lesson Plan"):
    # Create a new Word document
    document = docx.Document()

    # Add subject as a bold, navy-colored heading
    title = document.add_heading(level=0)
    run = title.add_run(subject.upper())
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 128)  # Navy color

    # Add lesson details in a table
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Pt(180)
    table.columns[1].width = Pt(300)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "LES TITEL:"
    hdr_cells[1].text = lesson_title
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "GRAAD:"
    hdr_cells[1].text = grade
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "VANAF:"
    hdr_cells[1].text = str(start_date)
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "TOT:"
    hdr_cells[1].text = str(end_date)

    # Add section headings and their contents
    document.add_paragraph()
    document.add_heading("LES DOELWITTE", level=1)
    document.add_paragraph(lesson_objective)
    
    document.add_heading("INLEIDING", level=1)
    document.add_paragraph(lesson_introduction)
    
    document.add_heading("LES AKTIWITEITE", level=1)
    document.add_paragraph(lesson_activities)
    
    document.add_heading("MATERIAAL BENODIG", level=1)
    document.add_paragraph(materials_needed)
    
    document.add_heading("HUISWERK", level=1)
    document.add_paragraph(homework)
    
    document.add_heading("NOTAS", level=1)
    document.add_paragraph(notes)
    
    # Add teacher information in a table
    document.add_paragraph()
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Pt(180)
    table.columns[1].width = Pt(300)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "VOORLETTERS:"
    hdr_cells[1].text = teacher_name
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "VAN:"
    hdr_cells[1].text = teacher_surname
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "HANDTEKENING:"
    hdr_cells[1].text = ""

    # Save document to BytesIO object
    with io.BytesIO() as output:
        document.save(output)
        output.seek(0)
        # Create download button
        st.download_button(
            label="Download Lesson Plan",
            data=output,
            file_name="lesson_plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success("Your lesson plan has been created. Click the download button to save the file.")
    
    # Add a fun element
    st.balloons()
