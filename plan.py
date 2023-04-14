import streamlit as st
import os
import io
import docx
from datetime import date

# Set page title and icon
st.set_page_config(page_title="Lesson Plan Creator", page_icon=":books:")

# Create input fields
st.title("HOЁRSKOOL SAUL DAMON")
st.subheader("LESBEPLANNER")
st.write("Vul asseblief die volgende velde in om 'n nuwe lesplan te skep:")
st.write("")
subject = st.text_input("VAK")
lesson_title = st.text_input("LES TITEL")
grade = st.selectbox("GRAAD", ["9", "10", "11", "12"])
st.subheader("KIES JOU BEGIN- EN EINDDATUM")
start_date = st.date_input("VANAF", value=date.today())
end_date = st.date_input("TOT", value=date.today())
lesson_objective = st.text_area("LES DOELWITTE")
lesson_introduction = st.text_area("INLEIDING")
lesson_activities = st.text_area("LES AKTIWITEITE")
materials_needed = st.text_area("MATERIAAL BENODIG")
homework = st.text_area("HUISWERK")
notes = st.text_area("NOTAS")
st.subheader("OPVOEDER BESONDERHEDE")
teacher_name = st.text_input("VOORLETTERS")
teacher_surname = st.text_input("VAN")
st.write("")
st.write("Created by Mr. A.R Visagie @ Saul Damon High School")

# Create save button
if st.button("Create Lesson Plan"):
    # Create a new Word document
    document = docx.Document()
    # Add input values to the document
    document.add_heading(subject.upper(), level=0)
    document.add_paragraph("")
    document.add_paragraph("LES TITEL: " + lesson_title)
    document.add_paragraph("GRAAD: " + grade)
    document.add_paragraph("FROM: " + str(start_date))
    document.add_paragraph("TO: " + str(end_date))
    document.add_paragraph("")
    document.add_heading("LES DOELWITTE", level=1)
    document.add_paragraph(lesson_objective)
    document.add_heading("INLEIDING", level=1)
    document.add_paragraph(lesson_introduction)
    document.add_heading("LES AKTIWITEITE", level=1)
    document.add_paragraph(lesson_activities)
    document.add_heading("MATERIAAL BENODIG", level=1)
    document.add_paragraph(materials_needed)
    document.add_heading("HUISWERK", level=1)
    document.add_paragraph(homework)
    document.add_heading("NOTAS", level=1)
    document.add_paragraph(notes)
    document.add_paragraph("")
    document.add_paragraph("VOORLETTERS: " + teacher_name)
    document.add_paragraph("VAN: " + teacher_surname)
    document.add_paragraph("HANDTEKENNG: ")

    # Save document to BytesIO object
    with io.BytesIO() as output:
        document.save(output)
        output.seek(0)
        # Create download button
        st.download_button(
            label="Download Lesson Plan",
            data=output,
            file_name="lesson_plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success("Your lesson plan has been created. Click the download button to save the file.")
